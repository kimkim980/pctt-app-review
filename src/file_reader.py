from __future__ import annotations
from pathlib import Path
import hashlib
import json
import os
import time
import pandas as pd
from docx import Document
from pypdf import PdfReader

MAX_CELL_CHARS = int(os.getenv("MAX_CELL_CHARS", "300"))
MAX_MARKDOWN_CHARS = int(os.getenv("MAX_MARKDOWN_CHARS", "60000"))
MAX_EXCEL_ROWS_FAST = int(os.getenv("MAX_EXCEL_ROWS_FAST", "300"))
MAX_EXCEL_ROWS_FULL = int(os.getenv("MAX_EXCEL_ROWS_FULL", "3000"))
CACHE_DIR = Path(os.getenv("PCTT_CACHE_DIR", ".cache"))


def df_to_markdown_safe(df: pd.DataFrame) -> str:
    """Convert DataFrame to markdown; fallback to plain table if tabulate is missing."""
    try:
        return df.to_markdown(index=False)
    except Exception:
        return df.to_string(index=False)


def _trim_cell(x):
    if pd.isna(x):
        return ""
    s = str(x).replace("\r", " ").replace("\n", " ").strip()
    return s[:MAX_CELL_CHARS] + ("..." if len(s) > MAX_CELL_CHARS else "")


def _cache_key(path: str, fast: bool, max_chars: int, max_rows: int) -> str:
    p = Path(path)
    stat = p.stat()
    raw = f"{p.resolve()}|{stat.st_mtime_ns}|{stat.st_size}|{fast}|{max_chars}|{max_rows}|v3"
    return hashlib.sha1(raw.encode("utf-8", errors="ignore")).hexdigest()


def _read_cache(path: str, fast: bool, max_chars: int, max_rows: int) -> str | None:
    try:
        CACHE_DIR.mkdir(parents=True, exist_ok=True)
        fp = CACHE_DIR / (_cache_key(path, fast, max_chars, max_rows) + ".md")
        if fp.exists():
            return fp.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return None
    return None


def _write_cache(path: str, fast: bool, max_chars: int, max_rows: int, content: str) -> None:
    try:
        CACHE_DIR.mkdir(parents=True, exist_ok=True)
        fp = CACHE_DIR / (_cache_key(path, fast, max_chars, max_rows) + ".md")
        fp.write_text(content, encoding="utf-8")
    except Exception:
        pass


def excel_to_dataframes(path: str, max_rows: int | None = None) -> dict[str, pd.DataFrame]:
    # nrows makes Markdown preview much faster for big CSDL files.
    return pd.read_excel(path, sheet_name=None, dtype=str, nrows=max_rows, engine=None)


def file_to_markdown(path: str, max_chars: int = MAX_MARKDOWN_CHARS, fast: bool = True) -> str:
    p = Path(path)
    suffix = p.suffix.lower()
    max_rows = MAX_EXCEL_ROWS_FAST if fast else MAX_EXCEL_ROWS_FULL

    cached = _read_cache(str(p), fast, max_chars, max_rows)
    if cached is not None:
        return cached

    parts: list[str] = []
    if suffix in [".xlsx", ".xls"]:
        sheets = excel_to_dataframes(str(p), max_rows=max_rows)
        for name, df in sheets.items():
            total_note = ""
            if len(df) >= max_rows:
                total_note = f"\n\n> Ghi chú: Fast preview chỉ chuyển {max_rows} dòng đầu của sheet này sang Markdown. Rule engine vẫn kiểm tra file Excel riêng."
            df = df.map(_trim_cell).fillna("")
            parts.append(f"## Sheet: {name}{total_note}\n\n" + df_to_markdown_safe(df))
    elif suffix == ".csv":
        df = pd.read_csv(p, dtype=str, nrows=max_rows).map(_trim_cell).fillna("")
        parts.append(df_to_markdown_safe(df))
    elif suffix == ".docx":
        doc = Document(str(p))
        paras = [x.text.strip() for x in doc.paragraphs if x.text.strip()]
        for i, table in enumerate(doc.tables, 1):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows[:max_rows]]
            if rows:
                if len(rows) > 1:
                    parts.append(f"## Table {i}\n" + df_to_markdown_safe(pd.DataFrame(rows[1:], columns=rows[0])))
                else:
                    parts.append(f"## Table {i}\n" + " | ".join(rows[0]))
        parts.insert(0, "\n\n".join(paras))
    elif suffix == ".pdf":
        reader = PdfReader(str(p))
        max_pages = 12 if fast else len(reader.pages)
        for i, page in enumerate(reader.pages[:max_pages], 1):
            parts.append(f"## Page {i}\n" + (page.extract_text() or ""))
        if fast and len(reader.pages) > max_pages:
            parts.append(f"...[Fast mode đã đọc {max_pages}/{len(reader.pages)} trang PDF]")
    elif suffix in [".md", ".txt"]:
        parts.append(p.read_text(encoding="utf-8", errors="ignore"))
    else:
        raise ValueError(f"Chưa hỗ trợ định dạng: {suffix}")
    md = "\n\n".join(parts)
    if len(md) > max_chars:
        md = md[:max_chars] + "\n\n...[Đã cắt bớt nội dung do quá dài. Bật Full mode nếu cần phân tích sâu hơn]"
    _write_cache(str(p), fast, max_chars, max_rows, md)
    return md
