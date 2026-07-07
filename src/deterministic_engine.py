from __future__ import annotations
from dataclasses import dataclass, asdict
from pathlib import Path
import pandas as pd
import re
import unicodedata
from rapidfuzz import process, fuzz
from .utils import normalize_text, truthy, to_float

@dataclass
class CheckItem:
    rule_group: str
    rule_name: str
    result: str
    severity: str
    evidence: str
    gap: str
    recommendation: str
    source_file: str = ""
    source_sheet: str = ""
    source_row: str = ""
    source_column: str = ""
    source_cell: str = ""
    source_value: str = ""
    abnormal_type: str = ""
    rule_source: str = ""

CANONICAL_ALIASES = {
    "site_code": ["ma tram", "mã trạm", "site", "site code", "ma site", "trạm", "ten tram", "tên trạm"],
    "priority": ["ut", "ưu tiên", "uu tien", "priority", "loai tram", "loại trạm"],
    "staff": ["em quan", "ém quân", "nhan su", "nhân sự", "ds em quan", "nguoi truc", "người trực"],
    "flood": ["ngap", "ngập", "nguy co ngap", "nguy cơ ngập", "lụt", "lut"],
    "cutoff": ["chia cat", "chia cắt", "co lap", "cô lập", "kho tiep can", "khó tiếp cận"],
    "ats": ["ats"],
    "generator_plan": ["pa cmn", "phuong an chay may no", "phương án chạy máy nổ", "mpd", "mpđ", "may no", "máy nổ"],
    "battery_hours": ["tgx", "thoi gian xa", "thời gian xả", "acquy", "ắc quy", "battery"],
    "distance_km": ["khoang cach", "khoảng cách", "km", "cự ly", "cu ly"],
    "access_time": ["thoi gian tiep can", "thời gian tiếp cận", "tiep can", "tiếp cận"],
}

def _norm_col(c: str) -> str:
    return normalize_text(c).lower()

def _excel_col_name(n: int) -> str:
    # n is 1-based
    out = ""
    while n:
        n, r = divmod(n - 1, 26)
        out = chr(65 + r) + out
    return out

def _cell_ref(df: pd.DataFrame, row_index: int, col_name: str | None) -> tuple[str, str, str]:
    if not col_name or col_name not in df.columns:
        return "", "", ""
    col_pos = list(df.columns).index(col_name) + 1
    excel_row = row_index + 2
    col_letter = _excel_col_name(col_pos)
    val = df.iloc[row_index, col_pos - 1]
    return col_letter, f"{col_letter}{excel_row}", "" if pd.isna(val) else str(val)

def _make_check(rule_group, rule_name, result, severity, evidence, gap, recommendation, fname, sheet, rowno='', df=None, row_index=None, col_name=None, abnormal_type='', rule_source='Rule engine CSDL'):
    source_column = source_cell = source_value = ""
    if df is not None and row_index is not None and col_name:
        source_column, source_cell, source_value = _cell_ref(df, int(row_index), col_name)
    return CheckItem(rule_group, rule_name, result, severity, evidence, gap, recommendation, fname, sheet, str(rowno), source_column, source_cell, source_value, abnormal_type, rule_source)

def infer_columns(df: pd.DataFrame) -> dict[str, str]:
    cols = list(map(str, df.columns))
    norm_map = {_norm_col(c): c for c in cols}
    result = {}
    for key, aliases in CANONICAL_ALIASES.items():
        candidates = []
        for nc, orig in norm_map.items():
            score = max(fuzz.partial_ratio(nc, a) for a in aliases)
            if score >= 78:
                candidates.append((orig, score))
        if candidates:
            result[key] = sorted(candidates, key=lambda x: x[1], reverse=True)[0][0]
    # Excel column-letter fallback after pandas preserves unnamed or actual headers poorly
    for letter_key, idx in {"staff":20, "cutoff":24, "flood":25, "ats":44, "battery_hours":41}.items():
        if letter_key not in result and len(cols) > idx:
            result[letter_key] = cols[idx]
    return result

def _iter_excel(path: str):
    sheets = pd.read_excel(path, sheet_name=None, dtype=str)
    for sheet, df in sheets.items():
        if len(df) == 0:
            continue
        yield sheet, df.fillna("")

def analyze_structured_files(paths: list[str]) -> dict:
    checks: list[CheckItem] = []
    for path in paths:
        p = Path(path)
        if p.suffix.lower() not in [".xlsx", ".xls", ".csv"]:
            continue
        if p.suffix.lower() == ".csv":
            sheets = [("CSV", pd.read_csv(p, dtype=str).fillna(""))]
        else:
            sheets = list(_iter_excel(str(p)))
        for sheet, df in sheets:
            cols = infer_columns(df)
            if not cols:
                checks.append(_make_check("CSDL BTS", "Nhận diện cột dữ liệu", "KHONG_DU_DU_LIEU", "MEDIUM", f"{p.name}/{sheet}: không nhận diện được cột chuẩn", "Thiếu header hoặc header không rõ", "Chuẩn hóa header hoặc mapping cột trong file CSDL", p.name, sheet, "", abnormal_type="Thiếu mapping/header"))
                continue
            checks.extend(_check_sheet(p.name, sheet, df, cols))
    score = _score(checks)
    return {
        "summary": {"deterministic_score": score, "deterministic_checks": len(checks), "deterministic_failed": sum(c.result in ["KHONG_DAT", "CAN_BO_SUNG"] for c in checks)},
        "checks": [asdict(c) for c in checks]
    }

def _row_id(row, cols):
    c = cols.get("site_code")
    return normalize_text(row.get(c, "")) if c else ""

def _check_sheet(fname, sheet, df, cols) -> list[CheckItem]:
    out: list[CheckItem] = []
    required = ["site_code", "priority", "staff", "flood", "cutoff", "ats", "generator_plan", "battery_hours"]
    missing = [x for x in required if x not in cols]
    if missing:
        out.append(_make_check("CSDL BTS", "Đủ trường dữ liệu tối thiểu", "CAN_BO_SUNG", "MEDIUM", f"Nhận diện được: {cols}", f"Thiếu/không nhận diện: {', '.join(missing)}", "Bổ sung mapping cột hoặc đổi header theo mẫu rule", fname, sheet, "", abnormal_type="Thiếu trường dữ liệu"))
    for i, row in df.iterrows():
        site = _row_id(row, cols) or f"row {i+2}"
        rowno = str(i + 2)
        priority = normalize_text(row.get(cols.get("priority", ""), "")).upper()
        staff = row.get(cols.get("staff", ""), "")
        flood = row.get(cols.get("flood", ""), "")
        cutoff = row.get(cols.get("cutoff", ""), "")
        ats = row.get(cols.get("ats", ""), "")
        plan = normalize_text(row.get(cols.get("generator_plan", ""), ""))
        batt = to_float(row.get(cols.get("battery_hours", ""), ""), None)
        dist = to_float(row.get(cols.get("distance_km", ""), ""), None)
        access = to_float(row.get(cols.get("access_time", ""), ""), None)
        if ("UT1" in priority) and not truthy(staff):
            out.append(_make_check("CSDL BTS", "100% trạm UT1/UT1_3321 phải có nhân sự ém quân", "KHONG_DAT", "HIGH", f"{site}: priority={priority}, nhân sự ém quân trống", "Thiếu nhân sự ém quân", "Bổ sung họ tên, SĐT, vị trí ém quân và ca trực", fname, sheet, rowno, df, i, cols.get("staff"), "Ô trống/bất thường"))
        if (truthy(flood) or truthy(cutoff)) and not truthy(staff):
            out.append(_make_check("CSDL BTS", "Trạm ngập/chia cắt phải có phương án ém quân", "KHONG_DAT", "CRITICAL", f"{site}: ngập={flood}, chia cắt={cutoff}, nhân sự trống", "Không có lực lượng tại chỗ khi thiên tai", "Bố trí ém quân hoặc nêu rõ PA tiếp cận thay thế có MPĐ dầu + ATS", fname, sheet, rowno, df, i, cols.get("staff"), "Thiếu PA ém quân"))
        if truthy(ats) and plan in {"3", "4", "5", "6"}:
            out.append(_make_check("CSDL BTS", "Logic ATS với phương án chạy máy nổ", "CAN_BO_SUNG", "MEDIUM", f"{site}: ATS={ats}, PA CMN={plan}", "Trạm có ATS nhưng PA CMN thuộc nhóm chạy lưu động/tòa nhà/BKK", "Rà soát lại PA đặt máy nổ; nếu vẫn chọn 3-6 cần ghi rõ căn cứ", fname, sheet, rowno, df, i, cols.get("generator_plan"), "Sai logic ATS/MPĐ"))
        if batt is not None and batt == 0:
            out.append(_make_check("CSDL BTS", "Thời gian xả ắc quy không được bất thường bằng 0", "CAN_BO_SUNG", "MEDIUM", f"{site}: TGX={batt}", "TGX = 0 có khả năng thiếu hoặc sai dữ liệu", "Kiểm tra lại PMCĐ/IMES và cập nhật TGX sau lắp đặt ắc quy", fname, sheet, rowno, df, i, cols.get("battery_hours"), "Giá trị bằng 0"))
        if batt is not None and batt < 2 and ((dist is not None and dist > 30) or (access is not None and access > batt)):
            out.append(_make_check("CSDL BTS", "TGX phải tương quan với thời gian/khoảng cách tiếp cận MPĐ", "KHONG_DAT", "HIGH", f"{site}: TGX={batt}h, khoảng cách={dist}, thời gian tiếp cận={access}", "TGX thấp trong khi tiếp cận xa/khó", "Đổi điểm ém quân, tăng nguồn dự phòng hoặc bố trí MPĐ/nhân sự gần trạm", fname, sheet, rowno, df, i, cols.get("battery_hours"), "Tương quan TGX/tiếp cận bất thường"))
    if not out:
        out.append(_make_check("CSDL BTS", "Rule engine CSDL", "DAT", "LOW", f"{fname}/{sheet}: không phát hiện lỗi cứng trên {len(df)} dòng", "", "Duy trì rà soát bằng AI đối với nội dung thuyết minh và phụ lục", fname, sheet, "", abnormal_type="Không phát hiện"))
    return out

def _score(checks: list[CheckItem]) -> int:
    score = 100
    penalty = {"CRITICAL": 18, "HIGH": 10, "MEDIUM": 5, "LOW": 2}
    for c in checks:
        if c.result == "KHONG_DAT":
            score -= penalty.get(c.severity, 5)
        elif c.result == "CAN_BO_SUNG":
            score -= max(2, penalty.get(c.severity, 4)//2)
    return max(0, min(100, score))


# ===== Text/phuong-an rule engine =====
def _read_docx_blocks(path: str) -> list[dict]:
    """Extract paragraphs and tables with rough positions for reporting."""
    try:
        from docx import Document
    except Exception:
        return []
    blocks = []
    doc = Document(path)
    for idx, para in enumerate(doc.paragraphs, 1):
        text = (para.text or "").strip()
        if text:
            blocks.append({"kind": "Đoạn", "pos": str(idx), "cell": f"Đoạn {idx}", "text": text})
    for ti, table in enumerate(doc.tables, 1):
        for ri, row in enumerate(table.rows, 1):
            cells = []
            for ci, cell in enumerate(row.cells, 1):
                txt = " ".join((cell.text or "").split())
                if txt:
                    cells.append(txt)
            if cells:
                blocks.append({"kind": "Bảng", "pos": f"{ti}.{ri}", "cell": f"Bảng {ti} dòng {ri}", "text": " | ".join(cells)})
    return blocks


def _read_text_file_blocks(path: str) -> list[dict]:
    p = Path(path)
    try:
        text = p.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return []
    blocks = []
    for i, line in enumerate(text.splitlines(), 1):
        line = line.strip()
        if line:
            blocks.append({"kind": "Dòng", "pos": str(i), "cell": f"Dòng {i}", "text": line})
    return blocks


def _input_text_blocks(paths: list[str]) -> list[dict]:
    all_blocks = []
    for path in paths:
        p = Path(path)
        if p.suffix.lower() == ".docx":
            blocks = _read_docx_blocks(str(p))
        elif p.suffix.lower() in {".txt", ".md"}:
            blocks = _read_text_file_blocks(str(p))
        else:
            continue
        for b in blocks:
            b["file"] = p.name
        all_blocks.extend(blocks)
    return all_blocks




def _fold_text(value) -> str:
    s = normalize_text(value).lower()
    return unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode("ascii")

_STOP_TERMS = {
    "phuong", "an", "dam", "bao", "trong", "truong", "hop", "cong", "tac", "noi", "dung", "tham", "dinh",
    "kiem", "tra", "danh", "gia", "tinh", "kha", "thi", "yeu", "cau", "thuc", "hien", "cap", "nhat",
    "tram", "cac", "cho", "theo", "voi", "danh", "sach", "hien", "trang", "nguon", "luc", "uctt", "pctt", "thuyet", "minh", "nhom", "dam", "bao", "thu", "phu", "con", "lai", "uu", "tien",
}


def _important_terms(text: str) -> list[str]:
    n = _fold_text(text)
    # Keep useful Vietnamese technical phrases first.
    phrases = [
        "ghe", "xuong", "xe cau", "xang dau", "nhien lieu", "mpd", "may no", "accu", "acquy", "ats",
        "noc", "cau truyen hinh", "jaber", "pm pctt", "nocpro", "nhan su", "so dien thoai", "sdt",
        "em quan", "lo trinh", "tuyen duong", "ngap lut", "chia cat", "co lap", "vat tu", "bugi", "avr",
        "cu de", "tu dien", "doi sua chua", "kho tap ket", "phuong tien", "xe", "can xang", "pmcd", "imes",
        "pa pctt", "pa cmn", "tgx", "giam tai", "thoi gian tiep can", "khoang cach", "doi dau noi", "doi tiep xang", "sch", "ut1_3321", "uu tien", "thu phu", "con lai",
    ]
    found = []
    for ph in phrases:
        if ph in n:
            found.append(ph)
    words = re.findall(r"[a-z0-9]{3,}", n)
    for w in words:
        if w not in _STOP_TERMS and not w.isdigit() and w not in found:
            found.append(w)
        if len(found) >= 12:
            break
    return found[:12]


def _requirement_lines(definition: str) -> list[str]:
    lines = []
    for raw in str(definition or "").splitlines():
        x = raw.strip(" -+•\t")
        if not x:
            continue
        # Keep actionable/checklist lines; ignore pure headings that are too generic.
        if len(normalize_text(x)) >= 12:
            lines.append(x)
    if not lines and definition:
        lines = [str(definition)[:300]]
    return lines[:10]


_GLOBAL_TEXT_CACHE = {}

def _get_global_text_folded(blocks: list[dict]) -> str:
    key = id(blocks)
    if key not in _GLOBAL_TEXT_CACHE:
        _GLOBAL_TEXT_CACHE[key] = " ".join(b.get("folded_text") or _fold_text(b.get("text", "")) for b in blocks)
    return _GLOBAL_TEXT_CACHE[key]


def _find_best_block(blocks: list[dict], query: str) -> tuple[dict | None, int]:
    q = _fold_text(query)
    if not q or not blocks:
        return None, 0
    best, best_score = None, 0
    terms = _important_terms(query)
    for b in blocks:
        t = b.get("folded_text")
        if t is None:
            t = _fold_text(b.get("text", ""))
            b["folded_text"] = t
        score = fuzz.partial_ratio(q, t)
        # Boost/penalize by concrete domain term overlap so generic headings like "PHƯƠNG ÁN" do not pass every rule.
        if terms:
            hit = sum(1 for term in terms if term in t)
            term_score = int(hit / max(len(terms), 1) * 100)
            if hit == 0:
                score = min(score, 35)
            else:
                score = max(score, term_score)
        if score > best_score:
            best, best_score = b, int(score)
    return best, best_score



def _nearby_text(blocks: list[dict], best: dict | None, window: int = 8) -> str:
    if not best or best not in blocks:
        return "\n".join(b.get("text", "") for b in blocks)
    idx = blocks.index(best)
    lo = max(0, idx - window)
    hi = min(len(blocks), idx + window + 1)
    return "\n".join(b.get("text", "") for b in blocks[lo:hi])


def _rule_rows_from_excel(path: str) -> list[dict]:
    rows = []
    try:
        sheets = pd.read_excel(path, sheet_name=None, dtype=str).items()
    except Exception:
        return rows
    for sheet, df in sheets:
        df = df.fillna("")
        headers = [normalize_text(str(c)).lower() for c in df.columns]
        for idx, r in df.iterrows():
            vals = [str(x).strip() for x in r.tolist()]
            if not any(vals):
                continue
            # Common uploaded format: STT | Ten phu luc/Truong thong tin | Noi dung tham dinh/cong viec | Huong dan
            rule_no = vals[0] if len(vals) > 0 else ""
            name = vals[1] if len(vals) > 1 else ""
            definition = "\n".join(v for v in vals[2:] if v)
            if not name and len(vals) > 2:
                name = vals[2]
            low_name = _fold_text(name)
            low_def = _fold_text(definition)
            if low_name in {"ten phu luc", "truong thong tin", "noi dung cong viec"} or low_name == "":
                continue
            if len(low_name + low_def) < 10:
                continue
            rows.append({"file": Path(path).name, "sheet": sheet, "excel_row": str(int(idx) + 2), "rule_no": rule_no, "rule_name": name, "definition": definition})
    return rows


def analyze_text_rule_files(input_paths: list[str], rule_files: list[str]) -> dict:
    """Evaluate DOCX/TXT/MD phương án against imported rule Excel files without relying on AI.

    This gives an offline, deterministic checklist: every row in a rule file becomes one rule item,
    with DAT/CAN_BO_SUNG/KHONG_DAT and the best matching document position.
    """
    blocks = _input_text_blocks(input_paths)
    checks: list[CheckItem] = []
    if not blocks:
        return {"summary": {"text_rule_checks": 0, "text_rule_failed": 0, "text_rule_score": 100}, "checks": []}
    all_text_norm = _fold_text("\n".join(b.get("text", "") for b in blocks))
    for rule_file in rule_files:
        for rr in _rule_rows_from_excel(rule_file):
            sheet_norm = _fold_text(rr.get("sheet", ""))
            # Prefer text/phuong-an sheets. Data sheets are handled by structured engine against xlsx/csv inputs.
            if any(k in sheet_norm for k in ["du lieu", "csdl", "data"]):
                continue
            name = rr.get("rule_name", "")
            definition = rr.get("definition", "")
            full_rule = (name + "\n" + definition).strip()
            best, score = _find_best_block(blocks, name or full_rule)
            nearby = _fold_text(_nearby_text(blocks, best))
            req_lines = _requirement_lines(definition)
            missing_reqs = []
            hit_reqs = []
            for req in req_lines:
                terms = _important_terms(req)
                # A requirement is considered covered if at least 45% of concrete terms are found near the best section or globally.
                if not terms:
                    continue
                hits = [t for t in terms if t in nearby]
                coverage = len(hits) / max(len(terms), 1)
                if coverage >= 0.45:
                    hit_reqs.append(req)
                else:
                    missing_reqs.append(req)
            result = "DAT"
            severity = "LOW"
            abnormal_type = "Đã có nội dung phù hợp"
            gap = ""
            recommendation = "Duy trì nội dung và bổ sung phụ lục/bằng chứng nếu có cập nhật."
            if score < 50 and not hit_reqs:
                result = "KHONG_DAT"
                severity = "HIGH"
                abnormal_type = "Không tìm thấy mục/phụ lục theo rule"
                gap = "Không tìm thấy nội dung tương ứng trong file phương án."
                recommendation = "Bổ sung mục/phụ lục theo rule, có số liệu, đầu mối chịu trách nhiệm và mốc hoàn thành."
            elif not req_lines and score < 90:
                result = "KHONG_DU_DU_LIEU"
                severity = "MEDIUM"
                abnormal_type = "Rule chưa có đủ căn cứ đối chiếu"
                gap = "Rule không có checklist chi tiết hoặc không tìm thấy mục tương ứng rõ trong phương án."
                recommendation = "Bổ sung nội dung thuyết minh/phụ lục tương ứng, hoặc cập nhật rule có tiêu chí kiểm tra cụ thể."
            elif missing_reqs:
                result = "CAN_BO_SUNG"
                severity = "MEDIUM"
                abnormal_type = "Thiếu nội dung/bằng chứng theo checklist"
                gap = "Thiếu: " + "; ".join(missing_reqs[:3])
                recommendation = "Bổ sung rõ các ý còn thiếu: tên đơn vị/cá nhân phụ trách, SĐT, số lượng, vị trí, thời gian hoàn thành và phụ lục chứng minh."
            evidence = ""
            if best:
                txt = best.get("text", "")
                evidence = (txt[:500] + "...") if len(txt) > 500 else txt
            else:
                evidence = "Không tìm thấy đoạn phù hợp."
            fname = best.get("file", Path(input_paths[0]).name) if best else Path(input_paths[0]).name
            source_row = best.get("cell", "") if best else ""
            checks.append(CheckItem(
                rule_group=f"Phương án/{rr.get('sheet','')}",
                rule_name=f"{rr.get('rule_no','')} - {name}".strip(" -"),
                result=result,
                severity=severity,
                evidence=f"Rule file {rr.get('file')} sheet {rr.get('sheet')} dòng {rr.get('excel_row')}. Bằng chứng gần nhất: {evidence}",
                gap=gap,
                recommendation=recommendation,
                source_file=fname,
                source_sheet="DOCX/TEXT",
                source_row=source_row,
                source_column="",
                source_cell=source_row,
                source_value=evidence,
                abnormal_type=abnormal_type,
                rule_source=f"Rule thuyết minh: {rr.get('file')}!{rr.get('sheet')}!A{rr.get('excel_row')}",
            ))
    failed = sum(c.result in {"KHONG_DAT", "CAN_BO_SUNG", "KHONG_DU_DU_LIEU"} for c in checks)
    score = _score(checks) if checks else 100
    return {"summary": {"text_rule_checks": len(checks), "text_rule_failed": failed, "text_rule_score": score}, "checks": [asdict(c) for c in checks]}


def analyze_all_files(input_paths: list[str], rule_files: list[str] | None = None) -> dict:
    """Combined deterministic engine for CSDL files and phương án text/docx files."""
    structured = analyze_structured_files(input_paths)
    text = analyze_text_rule_files(input_paths, rule_files or []) if rule_files else {"summary": {"text_rule_checks": 0, "text_rule_failed": 0, "text_rule_score": 100}, "checks": []}
    checks = structured.get("checks", []) + text.get("checks", [])
    scores = []
    if structured.get("summary", {}).get("deterministic_checks", 0):
        scores.append(int(structured.get("summary", {}).get("deterministic_score", 100)))
    if text.get("summary", {}).get("text_rule_checks", 0):
        scores.append(int(text.get("summary", {}).get("text_rule_score", 100)))
    score = min(scores) if scores else 100
    base_summary = {}
    base_summary.update(structured.get("summary", {}))
    base_summary.update(text.get("summary", {}))
    base_summary.update({
        "deterministic_score": score,
        "deterministic_checks": len(checks),
        "deterministic_failed": sum(c.get("result") in ["KHONG_DAT", "CAN_BO_SUNG", "KHONG_DU_DU_LIEU"] for c in checks),
    })
    return {"summary": base_summary, "checks": checks}

# ===== Rule-all-files engine v2: every uploaded rule file becomes checklist items =====
# These definitions intentionally override the earlier text-rule helpers above.

def _read_pdf_blocks(path: str) -> list[dict]:
    blocks = []
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        for pi, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            for li, line in enumerate(text.splitlines(), 1):
                line = " ".join(line.split())
                if line:
                    blocks.append({"kind": "Trang", "pos": f"{pi}.{li}", "cell": f"Trang {pi} dòng {li}", "text": line})
    except Exception:
        pass
    return blocks


def _read_excel_blocks(path: str, max_rows_per_sheet: int | None = None) -> list[dict]:
    blocks = []
    try:
        sheets = pd.read_excel(path, sheet_name=None, dtype=str)
    except Exception:
        return blocks
    for sheet, df in sheets.items():
        df = df.fillna("")
        if max_rows_per_sheet:
            df = df.head(max_rows_per_sheet)
        headers = [str(c) for c in df.columns]
        # Header block helps match rules that require specific fields.
        if headers:
            blocks.append({"kind": "Sheet", "pos": sheet, "cell": f"Sheet {sheet} header", "sheet": sheet, "text": " | ".join(headers)})
        for ri, row in df.iterrows():
            vals = []
            for ci, col in enumerate(df.columns, 1):
                val = str(row.get(col, "")).strip()
                if val:
                    vals.append(f"{col}: {val}")
            if vals:
                excel_row = int(ri) + 2
                blocks.append({"kind": "Excel", "pos": str(excel_row), "cell": f"Sheet {sheet} dòng {excel_row}", "sheet": sheet, "text": " | ".join(vals)})
    return blocks


def _input_text_blocks(paths: list[str]) -> list[dict]:
    """Read all candidate files into searchable blocks with positions."""
    all_blocks = []
    for path in paths:
        p = Path(path)
        suffix = p.suffix.lower()
        if suffix == ".docx":
            blocks = _read_docx_blocks(str(p))
        elif suffix == ".pdf":
            blocks = _read_pdf_blocks(str(p))
        elif suffix in {".txt", ".md"}:
            blocks = _read_text_file_blocks(str(p))
        elif suffix in {".xlsx", ".xls", ".csv"}:
            if suffix == ".csv":
                try:
                    df = pd.read_csv(p, dtype=str, nrows=200).fillna("")
                    blocks = []
                    for i, row in df.iterrows():
                        vals = [f"{c}: {row.get(c, '')}" for c in df.columns if str(row.get(c, '')).strip()]
                        if vals:
                            blocks.append({"kind": "CSV", "pos": str(i+2), "cell": f"Dòng {i+2}", "sheet": "CSV", "text": " | ".join(vals)})
                except Exception:
                    blocks = []
            else:
                blocks = _read_excel_blocks(str(p), max_rows_per_sheet=200)
        else:
            blocks = []
        for b in blocks:
            b["file"] = p.name
            if "sheet" not in b:
                b["sheet"] = b.get("kind", "TEXT")
            # Pre-calculate folded_text to speed up calculations.
            b["folded_text"] = _fold_text(b.get("text", ""))
        all_blocks.extend(blocks)
    return all_blocks



def _split_rule_text_to_items(text: str, source_file: str, source_sheet: str = "TEXT", base_row: str = "") -> list[dict]:
    """Convert a free-form rule document/text into individual rule items."""
    items = []
    current_title = ""
    lines = [" ".join(x.strip().split()) for x in str(text or "").splitlines()]
    counter = 0
    for line in lines:
        if not line:
            continue
        folded = _fold_text(line)
        # Identify headings/checklist bullets/numbered clauses.
        is_heading = bool(re.match(r"^([ivx]+\.|\d+(\.\d+)*\.?|[a-z]\)|[-+*•])\s+", line.lower())) or line.isupper()
        if len(folded) < 8:
            continue
        if is_heading and len(line) <= 160:
            current_title = line.strip("-+*• ")
        # Keep actionable lines as rules; avoid pure legal/form headings that are too generic.
        terms = _important_terms(line)
        if len(line) >= 20 and (terms or any(k in folded for k in ["phai", "can", "yeu cau", "dam bao", "kiem tra", "bo sung", "danh sach", "phu luc"])):
            counter += 1
            name = current_title if current_title and current_title != line else line[:120]
            items.append({
                "file": source_file,
                "sheet": source_sheet,
                "excel_row": base_row or str(counter),
                "rule_no": str(counter),
                "rule_name": name,
                "definition": line,
            })
    return items


def _rule_rows_from_excel(path: str) -> list[dict]:
    """Extract EVERY meaningful row from every sheet of an Excel rule file.

    No sheet is skipped. This is important because user-imported rule files can be CSDL, thuyet minh,
    phuong an, or custom checklists, and the tool must evaluate against all assigned rule files.
    """
    rows = []
    try:
        sheets = pd.read_excel(path, sheet_name=None, dtype=str)
    except Exception:
        return rows
    for sheet, df in sheets.items():
        df = df.fillna("")
        columns = [str(c).strip() for c in df.columns]
        folded_cols = [_fold_text(c) for c in columns]
        for idx, r in df.iterrows():
            vals = [str(x).strip() for x in r.tolist()]
            if not any(vals):
                continue
            # Skip repeated header rows only.
            folded_vals = [_fold_text(v) for v in vals]
            if set(folded_vals[: min(len(folded_vals), len(folded_cols))]) == set(folded_cols[: len(folded_vals)]):
                continue
            labeled_parts = []
            for c, v in zip(columns, vals):
                if v:
                    labeled_parts.append(f"{c}: {v}")
            # Pick a readable rule name from the most likely descriptive columns.
            name = ""
            for col_hint in ["noi dung", "tieu chi", "yeu cau", "hang muc", "muc", "truong thong tin", "ten phu luc", "rule"]:
                for c, v in zip(columns, vals):
                    if col_hint in _fold_text(c) and v:
                        name = v
                        break
                if name:
                    break
            if not name:
                # fallback: first non-empty non-STT value
                for v in vals[1:] + vals[:1]:
                    if v and len(_fold_text(v)) >= 3:
                        name = v
                        break
            definition = "\n".join(labeled_parts) if labeled_parts else "\n".join(v for v in vals if v)
            if len(_fold_text(name + " " + definition)) < 10:
                continue
            rows.append({
                "file": Path(path).name,
                "sheet": sheet,
                "excel_row": str(int(idx) + 2),
                "rule_no": vals[0] if vals else str(int(idx) + 2),
                "rule_name": name[:250],
                "definition": definition[:3000],
            })
    return rows


def _rule_rows_from_any_file(path: str) -> list[dict]:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix in {".xlsx", ".xls"}:
        return _rule_rows_from_excel(str(p))
    if suffix == ".csv":
        try:
            df = pd.read_csv(p, dtype=str).fillna("")
            tmp = p.with_suffix(".tmp.xlsx")
            # Avoid writing temp; parse DataFrame similarly.
            rows = []
            columns = [str(c).strip() for c in df.columns]
            for idx, r in df.iterrows():
                vals = [str(x).strip() for x in r.tolist()]
                if not any(vals):
                    continue
                labeled = [f"{c}: {v}" for c, v in zip(columns, vals) if v]
                name = next((v for v in vals if v), f"CSV dòng {idx+2}")
                rows.append({"file": p.name, "sheet": "CSV", "excel_row": str(int(idx)+2), "rule_no": vals[0] if vals else "", "rule_name": name[:250], "definition": "\n".join(labeled)[:3000]})
            return rows
        except Exception:
            return []
    if suffix == ".docx":
        blocks = _read_docx_blocks(str(p))
        text = "\n".join(b.get("text", "") for b in blocks)
        return _split_rule_text_to_items(text, p.name, "DOCX")
    if suffix == ".pdf":
        blocks = _read_pdf_blocks(str(p))
        text = "\n".join(b.get("text", "") for b in blocks)
        return _split_rule_text_to_items(text, p.name, "PDF")
    if suffix in {".txt", ".md"}:
        try:
            text = p.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            text = ""
        return _split_rule_text_to_items(text, p.name, suffix.upper().strip("."))
    return []


def _coverage_against_blocks(blocks: list[dict], rule_text: str, best: dict | None) -> tuple[float, list[str], list[str]]:
    terms = _important_terms(rule_text)
    if not terms:
        return 0.0, [], []
    global_text = _get_global_text_folded(blocks)
    if not best or best not in blocks:
        nearby = global_text
    else:
        idx = blocks.index(best)
        lo = max(0, idx - 10)
        hi = min(len(blocks), idx + 10 + 1)
        nearby = " ".join(b.get("folded_text") or _fold_text(b.get("text", "")) for b in blocks[lo:hi])
    hit = []
    missing = []
    for t in terms:
        if t in nearby or t in global_text:
            hit.append(t)
        else:
            missing.append(t)
    return len(hit) / max(len(terms), 1), hit, missing



def analyze_text_rule_files(input_paths: list[str], rule_files: list[str]) -> dict:
    """Evaluate every imported rule file against every imported thuyet-minh/phuong-an/CSDL file.

    Each meaningful row/line in each rule file becomes one checklist item. The output records the
    best matching evidence block and its position. This works offline and does not depend on GPT.
    """
    _GLOBAL_TEXT_CACHE.clear()
    blocks = _input_text_blocks(input_paths)

    checks: list[CheckItem] = []
    all_rules = []
    for rf in rule_files or []:
        all_rules.extend(_rule_rows_from_any_file(rf))
    if not blocks:
        checks.append(CheckItem(
            rule_group="Đầu vào",
            rule_name="Đọc nội dung file cần thẩm định",
            result="KHONG_DU_DU_LIEU",
            severity="HIGH",
            evidence="Không đọc được nội dung từ các file cần thẩm định.",
            gap="File có thể rỗng, sai định dạng, bị khóa, hoặc PDF scan ảnh chưa OCR.",
            recommendation="Dùng file DOCX/XLSX/PDF text; nếu PDF scan cần chuyển OCR trước khi thẩm định.",
            abnormal_type="Không đọc được file đầu vào",
            rule_source="System",
        ))
        return {"summary": {"text_rule_checks": 1, "text_rule_failed": 1, "text_rule_score": 0}, "checks": [asdict(c) for c in checks]}
    if not all_rules:
        checks.append(CheckItem(
            rule_group="Rule",
            rule_name="Đọc toàn bộ file rule được gán",
            result="KHONG_DU_DU_LIEU",
            severity="HIGH",
            evidence="Không trích xuất được tiêu chí nào từ file rule đã chọn.",
            gap="Rule có thể đang ở dạng ảnh/format quá đặc thù hoặc file bị khóa.",
            recommendation="Chuẩn hóa rule dạng Excel có các cột STT, Nhóm, Tiêu chí/Nội dung thẩm định, Yêu cầu/Bằng chứng; hoặc dùng DOCX/TXT có bullet rõ ràng.",
            abnormal_type="Không đọc được rule",
            rule_source="System",
        ))
        return {"summary": {"text_rule_checks": 1, "text_rule_failed": 1, "text_rule_score": 0}, "checks": [asdict(c) for c in checks]}

    for rr in all_rules:
        name = rr.get("rule_name", "")
        definition = rr.get("definition", "")
        full_rule = (name + "\n" + definition).strip()
        # Match by full rule; this prevents generic names like "Phụ lục 1" from passing too easily.
        best, fuzzy_score = _find_best_block(blocks, full_rule)
        coverage, hit_terms, missing_terms = _coverage_against_blocks(blocks, full_rule, best)
        req_lines = _requirement_lines(definition)
        missing_reqs = []
        for req in req_lines[:8]:
            cov, _, _ = _coverage_against_blocks(blocks, req, best)
            if cov < 0.45:
                missing_reqs.append(req)

        result = "DAT"
        severity = "LOW"
        abnormal_type = "Đạt theo rule"
        gap = ""
        recommendation = "Duy trì nội dung/bằng chứng theo rule."
        # Strict scoring: no concrete term evidence = fail; partial evidence = need supplement.
        if fuzzy_score < 45 and coverage < 0.35:
            result = "KHONG_DAT"
            severity = "HIGH"
            abnormal_type = "Không tìm thấy nội dung tương ứng với rule"
            gap = "Không tìm thấy bằng chứng đủ rõ trong các file thẩm định. Thiếu từ khóa/chủ đề: " + ", ".join(missing_terms[:8])
            recommendation = "Bổ sung đúng mục theo rule, có số liệu/phụ lục/bảng minh chứng và đầu mối chịu trách nhiệm."
        elif coverage < 0.55 or missing_reqs:
            result = "CAN_BO_SUNG"
            severity = "MEDIUM"
            abnormal_type = "Có nhắc đến nhưng thiếu bằng chứng/chi tiết"
            detail = []
            if missing_terms:
                detail.append("thiếu/chưa rõ: " + ", ".join(missing_terms[:8]))
            if missing_reqs:
                detail.append("thiếu yêu cầu: " + "; ".join(missing_reqs[:3]))
            gap = "; ".join(detail) or "Nội dung mới khớp một phần với rule."
            recommendation = "Bổ sung thông tin định lượng, bảng/phụ lục, vị trí trạm, nhân sự, SĐT, thời gian, phương tiện hoặc căn cứ theo đúng rule."

        evidence = "Không tìm thấy đoạn phù hợp."
        fname = Path(input_paths[0]).name if input_paths else ""
        source_sheet = "TEXT"
        source_cell = ""
        if best:
            txt = best.get("text", "")
            evidence = (txt[:700] + "...") if len(txt) > 700 else txt
            fname = best.get("file", fname)
            source_sheet = best.get("sheet", best.get("kind", "TEXT"))
            source_cell = best.get("cell", "")
        checks.append(CheckItem(
            rule_group=f"Rule tổng hợp/{rr.get('sheet','')}",
            rule_name=f"{rr.get('rule_no','')} - {name}".strip(" -"),
            result=result,
            severity=severity,
            evidence=f"Rule {rr.get('file')}!{rr.get('sheet')}!dòng {rr.get('excel_row')}. Khớp {fuzzy_score}%, coverage {int(coverage*100)}%, từ khóa khớp: {', '.join(hit_terms[:8])}. Bằng chứng gần nhất: {evidence}",
            gap=gap,
            recommendation=recommendation,
            source_file=fname,
            source_sheet=source_sheet,
            source_row=source_cell,
            source_column="",
            source_cell=source_cell,
            source_value=evidence,
            abnormal_type=abnormal_type,
            rule_source=f"{rr.get('file')}!{rr.get('sheet')}!dòng {rr.get('excel_row')}",
        ))
    failed = sum(c.result in {"KHONG_DAT", "CAN_BO_SUNG", "KHONG_DU_DU_LIEU"} for c in checks)
    score = _score(checks) if checks else 100
    return {"summary": {"text_rule_checks": len(checks), "text_rule_failed": failed, "text_rule_score": score, "rule_files_loaded": len(rule_files or []), "rule_items_loaded": len(all_rules)}, "checks": [asdict(c) for c in checks]}
