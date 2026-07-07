from __future__ import annotations
from pathlib import Path
import json
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from .agent import build_markdown_report
from .utils import now_stamp, safe_name

DETAIL_COLUMNS = [
    "rule_source", "rule_group", "rule_name", "result", "severity", "abnormal_type",
    "source_file", "source_sheet", "source_row", "source_column", "source_cell", "source_value",
    "evidence", "gap", "recommendation",
]

def _checks_df(result: dict) -> pd.DataFrame:
    df = pd.DataFrame(result.get("checks", []))
    for col in DETAIL_COLUMNS:
        if col not in df.columns:
            df[col] = ""
    return df[DETAIL_COLUMNS]

def _rule_assessment_df(result: dict) -> pd.DataFrame:
    checks = _checks_df(result)
    if checks.empty:
        return pd.DataFrame(columns=["rule_source", "rule_group", "rule_name", "status", "so_loi", "so_bat_thuong", "o_du_lieu_lien_quan", "khuyen_nghi"])
    bad_values = {"KHONG_DAT", "CAN_BO_SUNG", "KHONG_DU_DU_LIEU"}
    rows = []
    group_cols = ["rule_source", "rule_group", "rule_name"]
    for keys, g in checks.groupby(group_cols, dropna=False):
        status = "DAT"
        if (g["result"].astype(str) == "KHONG_DAT").any():
            status = "KHONG_DAT"
        elif (g["result"].astype(str).isin(["CAN_BO_SUNG", "KHONG_DU_DU_LIEU"])).any():
            status = "CAN_BO_SUNG"
        abnormal = g[g["result"].astype(str).isin(bad_values)]
        cells = sorted(set(x for x in abnormal.get("source_cell", pd.Series(dtype=str)).astype(str).tolist() if x))
        recs = [x for x in abnormal.get("recommendation", pd.Series(dtype=str)).astype(str).tolist() if x]
        rows.append({
            "rule_source": keys[0],
            "rule_group": keys[1],
            "rule_name": keys[2],
            "status": status,
            "so_loi": int((g["result"].astype(str) == "KHONG_DAT").sum()),
            "so_bat_thuong": int(len(abnormal)),
            "o_du_lieu_lien_quan": ", ".join(cells[:30]) + (" ..." if len(cells) > 30 else ""),
            "khuyen_nghi": "; ".join(dict.fromkeys(recs[:5])),
        })
    return pd.DataFrame(rows)

def _abnormal_df(result: dict) -> pd.DataFrame:
    checks = _checks_df(result)
    if checks.empty:
        return checks
    bad_values = {"KHONG_DAT", "CAN_BO_SUNG", "KHONG_DU_DU_LIEU"}
    return checks[checks["result"].astype(str).isin(bad_values)].copy()

def export_report(result: dict, out_dir: str, base_name: str = "bao_cao_tham_dinh") -> dict[str, str]:
    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)
    stem = f"{safe_name(base_name)}_{now_stamp()}"
    md_path = out / f"{stem}.md"
    json_path = out / f"{stem}.json"
    xlsx_path = out / f"{stem}.xlsx"
    docx_path = out / f"{stem}.docx"
    md = result.get("report_markdown") or build_markdown_report(result)
    md += "\n\n## Đánh giá theo từng mục rule\n\n"
    rule_df = _rule_assessment_df(result)
    if len(rule_df):
        try:
            md += rule_df.to_markdown(index=False)
        except Exception:
            md += rule_df.to_string(index=False)
    else:
        md += "Không có dữ liệu đánh giá rule."
    json_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    md_path.write_text(md, encoding="utf-8")
    _export_xlsx(result, xlsx_path)
    _export_docx(result, docx_path)
    return {"markdown": str(md_path), "json": str(json_path), "xlsx": str(xlsx_path), "docx": str(docx_path)}

def _write_df(writer, df: pd.DataFrame, sheet_name: str):
    df.to_excel(writer, sheet_name=sheet_name, index=False)
    wb = writer.book
    ws = writer.sheets[sheet_name]
    header_fmt = wb.add_format({"bold": True, "bg_color": "#D9EAF7", "border": 1, "text_wrap": True})
    bad_fmt = wb.add_format({"bg_color": "#F4CCCC"})
    warn_fmt = wb.add_format({"bg_color": "#FFF2CC"})
    ok_fmt = wb.add_format({"bg_color": "#D9EAD3"})
    ws.freeze_panes(1, 0)
    ws.autofilter(0, 0, max(len(df), 1), max(len(df.columns) - 1, 0))
    for col_num, value in enumerate(df.columns):
        ws.write(0, col_num, value, header_fmt)
        width = min(max(12, len(str(value)) + 3), 45)
        if len(df) > 0:
            width = min(max(width, int(df.iloc[:, col_num].astype(str).str.len().quantile(0.9)) + 2), 70)
        ws.set_column(col_num, col_num, width)
    for status_col in ["result", "status"]:
        if status_col in df.columns and len(df) > 0:
            c = df.columns.get_loc(status_col)
            ws.conditional_format(1, c, len(df), c, {"type": "text", "criteria": "containing", "value": "KHONG_DAT", "format": bad_fmt})
            ws.conditional_format(1, c, len(df), c, {"type": "text", "criteria": "containing", "value": "CAN_BO_SUNG", "format": warn_fmt})
            ws.conditional_format(1, c, len(df), c, {"type": "text", "criteria": "containing", "value": "KHONG_DU_DU_LIEU", "format": warn_fmt})
            ws.conditional_format(1, c, len(df), c, {"type": "text", "criteria": "containing", "value": "DAT", "format": ok_fmt})

def _export_xlsx(result: dict, path: Path):
    summary = pd.DataFrame([result.get("summary", {})])
    rule_df = _rule_assessment_df(result)
    abnormal = _abnormal_df(result)
    checks = _checks_df(result)
    with pd.ExcelWriter(path, engine="xlsxwriter") as writer:
        _write_df(writer, summary, "Tong hop")
        _write_df(writer, rule_df, "Danh gia tung muc")
        _write_df(writer, abnormal, "Diem bat thuong")
        _write_df(writer, checks, "Chi tiet tat ca rule")

def _export_docx(result: dict, path: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BÁO CÁO THẨM ĐỊNH PCTT/UCTT TRẠM BTS")
    run.bold = True
    run.font.size = Pt(15)
    summary = result.get("summary", {})
    doc.add_heading("1. Kết luận tổng hợp", level=1)
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    data = [("Điểm tổng", summary.get("overall_score", summary.get("deterministic_score", "N/A"))), ("Kết luận", summary.get("overall_result", "N/A")), ("Phát hiện chính", "; ".join(map(str, summary.get("key_findings", []) or [])))]
    for row, (k, v) in zip(table.rows, data):
        row.cells[0].text = str(k)
        row.cells[1].text = str(v)

    doc.add_heading("2. Đánh giá theo từng mục rule", level=1)
    rule_df = _rule_assessment_df(result)
    if not rule_df.empty:
        cols = ["rule_group", "rule_name", "status", "so_bat_thuong", "o_du_lieu_lien_quan", "khuyen_nghi"]
        t = doc.add_table(rows=1, cols=len(cols))
        t.style = "Table Grid"
        headers = ["Nhóm", "Mục rule", "Đạt/Không đạt", "Số bất thường", "Ô dữ liệu", "Khuyến nghị"]
        for i, c in enumerate(headers):
            t.rows[0].cells[i].text = c
        for _, r in rule_df.iterrows():
            cells = t.add_row().cells
            for i, c in enumerate(cols):
                cells[i].text = str(r.get(c, ""))
    else:
        doc.add_paragraph("Không có dữ liệu đánh giá theo rule.")

    doc.add_heading("3. Điểm bất thường và vị trí ô dữ liệu", level=1)
    abnormal = _abnormal_df(result)
    if not abnormal.empty:
        cols = ["rule_name", "result", "severity", "source_file", "source_sheet", "source_cell", "source_value", "gap", "recommendation"]
        t = doc.add_table(rows=1, cols=len(cols))
        t.style = "Table Grid"
        headers = ["Rule", "KQ", "Mức", "File", "Sheet", "Ô", "Giá trị", "Bất thường", "Khuyến nghị"]
        for i, c in enumerate(headers):
            t.rows[0].cells[i].text = c
        for _, r in abnormal.head(200).iterrows():
            cells = t.add_row().cells
            for i, c in enumerate(cols):
                cells[i].text = str(r.get(c, ""))
        if len(abnormal) > 200:
            doc.add_paragraph(f"Danh sách Word chỉ hiển thị 200/{len(abnormal)} dòng. Xem đầy đủ trong file Excel sheet 'Diem bat thuong'.")
    else:
        doc.add_paragraph("Không phát hiện điểm bất thường.")

    doc.add_heading("4. Hướng xử lý ưu tiên", level=1)
    for text in ["Xử lý lỗi CRITICAL/HIGH trước khi phê duyệt.", "Mở file Excel báo cáo, lọc sheet 'Diem bat thuong' theo cột source_cell để sửa đúng ô dữ liệu.", "Sau khi sửa dữ liệu, chạy lại thẩm định để xác nhận trạng thái ĐẠT/KHÔNG ĐẠT theo từng rule."]:
        doc.add_paragraph(text, style="List Number")
    doc.save(path)
